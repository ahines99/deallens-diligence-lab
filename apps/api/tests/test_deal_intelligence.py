"""Focused offline coverage for versioned data-room and evidence intelligence APIs."""
from __future__ import annotations

import hashlib
import io

import pytest
from docx import Document as WordDocument
from fastapi import FastAPI, Request
from fastapi.testclient import TestClient
from openpyxl import Workbook
from sqlalchemy import create_engine
from sqlalchemy.orm import Session
from sqlalchemy.pool import StaticPool

from src.db.base import Base
from src.db.session import get_session
from src.models.deal_intelligence import ClaimReviewEvent
from src.models.document import DocumentChunk as FilingChunk
from src.models.filing import Filing
from src.models.workspace import Workspace  # noqa: F401 - registers referenced table
from src.routers.deal_intelligence import router as intelligence_router
from src.routers.deal_workflow import router as workflow_router
from src.schemas.deal_intelligence import (
    CitedQARequest,
    ClaimReviewRequest,
    ComparisonRequest,
    DocumentTextCreate,
    EvaluationCase,
    EvaluationRequest,
    ExtractionRequest,
    QAFilters,
    SecFilingComparisonRequest,
)
from src.schemas.deal_workflow import (
    ActorContext,
    DealCreate,
    FundCreate,
    OrganizationCreate,
)
from src.schemas.identity import PrincipalContext
from src.services import deal_intelligence_service as service
from src.services import deal_workflow_service as workflow


@pytest.fixture()
def db():
    engine = create_engine("sqlite:///:memory:")
    Base.metadata.create_all(engine)
    with Session(engine, expire_on_commit=False) as session:
        yield session
    engine.dispose()


def _setup(db: Session, suffix: str = "intel"):
    creator = ActorContext(actor_id=f"lead-{suffix}", display_name="Deal Lead")
    organization = workflow.create_organization(
        db,
        OrganizationCreate(name=f"Intelligence Org {suffix}", slug=f"intelligence-{suffix}"),
        creator,
    )
    actor = creator.model_copy(update={"organization_id": organization.id})
    fund = workflow.create_fund(db, organization.id, FundCreate(name="Fund I"), actor)
    deal = workflow.create_deal(
        db,
        fund.id,
        DealCreate(
            code=f"DI-{suffix}",
            name=f"Project {suffix}",
            target_company="Evidence Target",
        ),
        actor,
    )
    return actor, organization, fund, deal


def test_text_ingestion_versions_hashes_exact_citations_and_abstention(db: Session):
    actor, _, _, deal = _setup(db, "text")
    first = service.ingest_text_document(
        db,
        deal.id,
        DocumentTextCreate(
            filename="../../QoE Notes.txt",
            title="Quality of Earnings",
            text="FY2025 EBITDA was $20 million.\n\nRevenue was $110 million.",
            document_metadata={"workstream": "financial", "confidential": True},
        ),
        actor,
    )
    assert first.filename == "QoE Notes.txt"
    assert first.version == 1
    assert first.sha256 == hashlib.sha256(first.raw_bytes).hexdigest()
    assert first.original_filename == "../../QoE Notes.txt"
    chunks = service.list_chunks(db, first.id, actor)
    assert [item.locator["paragraph"] for item in chunks] == [1, 2]

    second = service.ingest_text_document(
        db,
        deal.id,
        DocumentTextCreate(
            filename="QoE Notes.txt",
            text="FY2025 EBITDA was $22 million.\n\nRevenue was $115 million.",
            logical_document_id=first.logical_document_id,
            document_metadata={"workstream": "financial", "confidential": True},
        ),
        actor,
    )
    assert second.version == 2
    assert second.supersedes_document_id == first.id
    assert service.list_documents(db, deal.id, actor)[0].id == second.id
    assert [item.version for item in service.list_document_versions(
        db, deal.id, first.logical_document_id, actor
    )] == [1, 2]

    answer = service.answer_question(
        db,
        deal.id,
        CitedQARequest(
            question="What was FY2025 EBITDA?",
            filters=QAFilters(document_ids=[first.id]),
        ),
        actor,
    )
    assert answer.status == "answered"
    assert answer.answer == "FY2025 EBITDA was $20 million."
    assert answer.citations[0]["quote"] == answer.answer
    assert answer.citations[0]["locator"] == {"type": "text", "paragraph": 1}
    assert answer.citations[0]["sha256"] == first.sha256

    abstention = service.answer_question(
        db,
        deal.id,
        CitedQARequest(
            question="How many employees are there?",
            filters=QAFilters(metadata={"workstream": "legal"}),
        ),
        actor,
    )
    assert abstention.status == "abstained"
    assert abstention.citations == []
    assert "could not find" in abstention.answer

    with pytest.raises(service.IntelligenceConflict, match="identical"):
        service.ingest_text_document(
            db,
            deal.id,
            DocumentTextCreate(
                filename="QoE Notes.txt",
                text="FY2025 EBITDA was $22 million.\n\nRevenue was $115 million.",
                logical_document_id=first.logical_document_id,
            ),
            actor,
        )


def test_compound_qa_preserves_decimals_and_cites_each_answer_sentence(db: Session):
    actor, _, _, deal = _setup(db, "compound")
    document = service.ingest_text_document(
        db,
        deal.id,
        DocumentTextCreate(
            filename="QoE Summary.txt",
            text=(
                "FY2025 revenue was $185.0 million and adjusted EBITDA was $38.25 million. "
                "The largest customer represented 14% of revenue in FY2024. "
                "First-lien debt bears interest at SOFR plus 450 basis points with a 1.0% floor."
            ),
        ),
        actor,
    )
    answer = service.answer_question(
        db,
        deal.id,
        CitedQARequest(
            question="What was adjusted EBITDA and what is the first-lien pricing?",
            filters=QAFilters(document_ids=[document.id]),
        ),
        actor,
    )
    assert answer.status == "answered"
    assert "$38.25 million" in answer.answer
    assert "450 basis points" in answer.answer
    assert "largest customer" not in answer.answer
    assert len(answer.citations) == 2
    assert all(citation["quote"] in answer.answer for citation in answer.citations)
    extracted = service.extract_structured_claims(
        db,
        deal.id,
        ExtractionRequest(document_ids=[document.id], categories=["kpi"]),
        actor,
    )
    assert any("$185.0 million" in claim.value_text for claim in extracted)


def test_docx_xlsx_csv_parsers_return_resolvable_source_locations(db: Session):
    actor, _, _, deal = _setup(db, "formats")

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Financials"
    sheet.append(["Metric", "FY2025"])
    sheet.append(["Revenue", 125_000_000])
    workbook_bytes = io.BytesIO()
    workbook.save(workbook_bytes)
    xlsx = service.ingest_document(
        db,
        deal.id,
        filename="Management Case.xlsx",
        content=workbook_bytes.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        actor=actor,
    )
    xlsx_chunks = service.list_chunks(db, xlsx.id, actor)
    assert xlsx_chunks[1].locator == {
        "type": "xlsx",
        "sheet": "Financials",
        "row": 2,
        "cell_start": "A2",
        "cell_end": "B2",
        "cell_range": "A2:B2",
    }
    assert "B2=125000000" in xlsx_chunks[1].text

    word = WordDocument()
    word.add_paragraph("The customer agreement renews annually.")
    word.add_paragraph("The largest customer represented 18% of FY2025 revenue.")
    word_bytes = io.BytesIO()
    word.save(word_bytes)
    docx = service.ingest_document(
        db,
        deal.id,
        filename="Commercial DD.docx",
        content=word_bytes.getvalue(),
        actor=actor,
    )
    docx_chunks = service.list_chunks(db, docx.id, actor)
    assert docx_chunks[0].locator == {"type": "docx", "paragraph": 1}
    assert docx_chunks[1].locator == {"type": "docx", "paragraph": 2}

    csv_document = service.ingest_text_document(
        db,
        deal.id,
        DocumentTextCreate(filename="cohorts.csv", text="Customer,ARR\nAlpha,2500000\n"),
        actor,
    )
    csv_chunks = service.list_chunks(db, csv_document.id, actor)
    assert csv_chunks[1].locator["cell_range"] == "A2:B2"

    with pytest.raises(service.IntelligenceError) as unsupported:
        service.ingest_document(
            db, deal.id, filename="malware.exe", content=b"not safe", actor=actor
        )
    assert unsupported.value.status_code == 415


def test_structured_extraction_review_history_and_approved_separation(db: Session):
    actor, _, _, deal = _setup(db, "claims")
    document = service.ingest_text_document(
        db,
        deal.id,
        DocumentTextCreate(
            filename="diligence.txt",
            text=(
                "FY2025 ARR was $120 million.\n"
                "The largest customer represented 22% of revenue.\n"
                "The term loan bears interest at SOFR plus 450 bps and matures in 2029.\n"
                "Management proposed a one-time $3 million add-back.\n"
                "The customer agreement renews annually and permits termination on 30 days notice."
            ),
        ),
        actor,
    )
    claims = service.extract_structured_claims(
        db,
        deal.id,
        ExtractionRequest(document_ids=[document.id]),
        actor,
    )
    categories = {claim.category for claim in claims}
    assert {"debt_term", "customer", "contract", "kpi", "qoe_candidate"} <= categories
    assert all(claim.review_status == "unreviewed" for claim in claims)
    assert all(claim.source_span["text"] in service.list_chunks(db, claim.document_id, actor)[
        next(
            index
            for index, chunk in enumerate(service.list_chunks(db, claim.document_id, actor))
            if chunk.id == claim.chunk_id
        )
    ].text for claim in claims)

    qoe = next(claim for claim in claims if claim.category == "qoe_candidate")
    reviewer = actor.model_copy(update={"actor_id": "qoe-reviewer"})
    approved, review = service.review_claim(
        db,
        qoe.id,
        ClaimReviewRequest(action="approve", expected_revision=1, note="QoE advisor confirmed"),
        reviewer,
    )
    assert approved.revision == 2
    assert approved.review_status == "approved"
    assert review.from_claim_id == qoe.id

    grouped = service.list_current_claims(db, deal.id, actor)
    assert [item.id for item in grouped["approved"]] == [approved.id]
    assert qoe.id not in {item.id for item in grouped["pending"]}
    revisions, reviews = service.claim_history(db, qoe.logical_claim_id, actor)
    assert [item.revision for item in revisions] == [1, 2]
    assert [item.action for item in reviews] == ["approve"]

    with pytest.raises(service.IntelligenceConflict, match="stale"):
        service.review_claim(
            db,
            qoe.id,
            ClaimReviewRequest(action="reject", expected_revision=1),
            reviewer,
        )

    review.note = "tampered"
    with pytest.raises(ValueError, match="append-only"):
        db.commit()
    db.rollback()
    assert db.get(ClaimReviewEvent, review.id).note == "QoE advisor confirmed"


def test_change_contradiction_and_guardrail_evaluation(db: Session):
    actor, _, _, deal = _setup(db, "compare")
    before = service.ingest_text_document(
        db,
        deal.id,
        DocumentTextCreate(
            filename="Debt Memo.txt",
            text="Total debt was $100 million. The agreement may be terminated.",
        ),
        actor,
    )
    after = service.ingest_text_document(
        db,
        deal.id,
        DocumentTextCreate(
            filename="Debt Memo.txt",
            text="Total debt was $140 million. The agreement may not be terminated.",
            logical_document_id=before.logical_document_id,
        ),
        actor,
    )
    change = service.compare_documents(
        db,
        deal.id,
        ComparisonRequest(
            from_document_id=before.id,
            to_document_id=after.id,
            comparison_type="change",
        ),
        actor,
    )
    assert change.finding_count == 1
    assert change.findings[0]["finding_type"] == "modified"
    assert change.findings[0]["before"]["sha256"] == before.sha256

    contradiction = service.compare_documents(
        db,
        deal.id,
        ComparisonRequest(
            from_document_id=before.id,
            to_document_id=after.id,
            comparison_type="contradiction",
        ),
        actor,
    )
    assert contradiction.finding_count >= 1
    assert contradiction.findings[0]["finding_type"] == "numeric_contradiction"

    evaluation = service.run_evaluation(
        db,
        deal.id,
        EvaluationRequest(
            cases=[
                EvaluationCase(
                    question="What was total debt?",
                    expected_answer_contains=["$140 million"],
                ),
                EvaluationCase(
                    question="What is the employee headcount?", should_abstain=True
                ),
            ]
        ),
        actor,
    )
    assert evaluation.passed is True
    assert evaluation.metrics["numeric_traceability"] == 1.0
    assert evaluation.metrics["citation_resolution"] == 1.0
    assert evaluation.metrics["abstention_accuracy"] == 1.0


def test_sec_filing_chunk_diff_has_accession_and_section_provenance(db: Session):
    workspace = Workspace(name="SEC Diff", deal_type="public_equity", status="draft")
    db.add(workspace)
    db.flush()
    before = Filing(
        workspace_id=workspace.id,
        company_name="Issuer",
        ticker="TEST",
        cik="1",
        form_type="10-K",
        filing_date="2024-12-31",
        accession_number="0001-24-000001",
        document_url="https://www.sec.gov/Archives/before.htm",
    )
    after = Filing(
        workspace_id=workspace.id,
        company_name="Issuer",
        ticker="TEST",
        cik="1",
        form_type="10-K",
        filing_date="2025-12-31",
        accession_number="0001-25-000001",
        document_url="https://www.sec.gov/Archives/after.htm",
    )
    db.add_all([before, after])
    db.flush()
    db.add_all(
        [
            FilingChunk(
                filing_id=before.id,
                workspace_id=workspace.id,
                section="Item 1A",
                chunk_index=0,
                chunk_text="Customer concentration was 15%.",
                source_url=before.document_url,
            ),
            FilingChunk(
                filing_id=after.id,
                workspace_id=workspace.id,
                section="Item 1A",
                chunk_index=0,
                chunk_text="Customer concentration was 28%.",
                source_url=after.document_url,
            ),
        ]
    )
    db.commit()

    comparison = service.compare_sec_filings(
        db,
        workspace.id,
        SecFilingComparisonRequest(from_filing_id=before.id, to_filing_id=after.id),
    )
    assert comparison.finding_count == 1
    finding = comparison.findings[0]
    assert finding["finding_type"] == "modified"
    assert finding["before"]["locator"]["section"] == "Item 1A"
    assert finding["after"]["locator"]["accession_number"] == "0001-25-000001"


def test_router_json_multipart_download_and_tenant_scope():
    engine = create_engine(
        "sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool
    )
    Base.metadata.create_all(engine)
    app = FastAPI()

    @app.middleware("http")
    async def verified_test_principal(request: Request, call_next):
        principal_organization = request.headers.get("X-Test-Principal-Organization")
        if principal_organization:
            request.state.principal = PrincipalContext(
                user_id="verified-outsider",
                session_id="verified-session",
                email="outsider@example.com",
                display_name="Verified Outsider",
                organization_id=principal_organization,
                membership_id="verified-membership",
                role="member",
            )
        return await call_next(request)

    app.include_router(workflow_router)
    app.include_router(intelligence_router)

    def session_override():
        with Session(engine, expire_on_commit=False) as session:
            yield session

    app.dependency_overrides[get_session] = session_override
    with TestClient(app) as client:
        organization = client.post(
            "/api/organizations", json={"name": "Router Org", "slug": "router-intel"}
        ).json()
        headers = {
            "X-Actor-ID": "router-lead",
            "X-Organization-ID": organization["id"],
        }
        fund = client.post(
            f"/api/organizations/{organization['id']}/funds",
            json={"name": "Router Fund"},
            headers=headers,
        ).json()
        deal = client.post(
            f"/api/funds/{fund['id']}/deals",
            json={"code": "RT-1", "name": "Router Deal", "target_company": "Target"},
            headers=headers,
        ).json()

        created = client.post(
            f"/api/deals/{deal['id']}/intelligence/documents",
            json={
                "filename": "management.txt",
                "text": "FY2025 revenue was $75 million.",
                "document_metadata": {"workstream": "financial"},
            },
            headers=headers,
        )
        assert created.status_code == 201, created.text
        document = created.json()

        upload = client.post(
            f"/api/deals/{deal['id']}/intelligence/documents/upload",
            files={"file": ("customers.csv", b"Customer,Revenue\nAlpha,50\n", "text/csv")},
            data={"metadata_json": '{"workstream":"commercial"}'},
            headers=headers,
        )
        assert upload.status_code == 201, upload.text
        assert upload.json()["extension"] == ".csv"

        qa = client.post(
            f"/api/deals/{deal['id']}/intelligence/qa",
            json={"question": "What was FY2025 revenue?"},
            headers=headers,
        )
        assert qa.status_code == 201
        assert qa.json()["citations"][0]["document_id"] == document["id"]

        download = client.get(
            f"/api/intelligence/documents/{document['id']}/download", headers=headers
        )
        assert download.status_code == 200
        assert download.content == b"FY2025 revenue was $75 million."
        assert "management.txt" in download.headers["content-disposition"]

        for suffix in ("", "/chunks", "/download"):
            hidden = client.get(
                f"/api/intelligence/documents/{document['id']}{suffix}",
                headers={
                    "X-Test-Principal-Organization": "f" * 32,
                    # Spoof the owning tenant in legacy headers; the verified principal wins.
                    "X-Organization-ID": organization["id"],
                    "X-Actor-ID": "spoofed-owner",
                },
            )
            assert hidden.status_code == 404
            assert hidden.json() == {"detail": "Document not found"}
    engine.dispose()
