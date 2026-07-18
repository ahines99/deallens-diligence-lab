"""Render frozen investment-committee packets into controlled, hashed export files."""
from __future__ import annotations

import hashlib
import io
import json
import re
from dataclasses import dataclass
from typing import Any
from xml.sax.saxutils import escape

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy import select
from sqlalchemy.orm import Session

from src.db.base import now_utc
from src.models.deal_workflow import (
    ConditionToClose,
    ICComment,
    ICDecision,
    ICPacketExport,
)
from src.schemas.deal_workflow import ActorContext, ExportRequest
from src.services import deal_workflow_service as workflow
from src.services import export_signing_service


@dataclass(frozen=True)
class ExportedICFile:
    content: bytes
    filename: str
    media_type: str
    sha256: str
    export_id: str


_MEDIA_TYPES = {
    "json": "application/json",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}


def _safe_json(value: Any) -> Any:
    return json.loads(json.dumps(value, default=lambda item: item.isoformat()))


def _file_stem(value: str) -> str:
    clean = re.sub(r"[^A-Za-z0-9._-]+", "-", value).strip("-.").lower()
    return clean[:80] or "deal"


def _flatten(value: Any, path: str = "") -> list[tuple[str, str]]:
    if isinstance(value, dict):
        rows: list[tuple[str, str]] = []
        for key in sorted(value):
            child = f"{path}.{key}" if path else str(key)
            rows.extend(_flatten(value[key], child))
        return rows
    if isinstance(value, list):
        rows = []
        for index, item in enumerate(value):
            rows.extend(_flatten(item, f"{path}[{index}]"))
        return rows or [(path, "[]")]
    return [(path, "" if value is None else str(value))]


def _excel_safe(value: Any) -> Any:
    if not isinstance(value, str):
        return value
    value = value[:32_767]
    if value.lstrip().startswith(("=", "+", "-", "@")):
        return "'" + value
    return value


def _packet_payload(session: Session, packet_id: str, actor: ActorContext | None) -> tuple[Any, Any, dict]:
    packet = workflow.get_ic_packet(session, packet_id, actor)
    if packet.status == "draft":
        raise workflow.WorkflowConflict("Freeze and submit an IC packet before exporting it")
    deal = workflow.get_deal(session, packet.deal_id, actor)
    decisions = list(
        session.scalars(
            select(ICDecision).where(ICDecision.packet_id == packet.id).order_by(ICDecision.sequence)
        )
    )
    comments = list(session.scalars(select(ICComment).where(ICComment.packet_id == packet.id)))
    conditions = list(
        session.scalars(select(ConditionToClose).where(ConditionToClose.packet_id == packet.id))
    )
    payload = {
        "export_metadata": {
            "generated_at": now_utc(),
            "organization_id": deal.organization_id,
            "fund_id": deal.fund_id,
            "deal_id": deal.id,
            "deal_code": deal.code,
            "deal_name": deal.name,
            "packet_id": packet.id,
            "packet_version": packet.version,
            "packet_status": packet.status,
            "packet_content_hash": packet.content_hash,
            "frozen_at": packet.frozen_at,
        },
        "decision_request": packet.decision_request,
        "scenario_snapshot": packet.scenario_snapshot,
        "model_snapshot": packet.model_snapshot,
        "thesis_snapshot": packet.thesis_snapshot,
        "risk_snapshot": packet.risk_snapshot,
        "evidence_manifest": packet.evidence_manifest,
        "review": {
            "comments": [
                {
                    "id": item.id,
                    "section_path": item.section_path,
                    "body": item.body,
                    "blocking": item.blocking,
                    "status": item.status,
                    "author_actor_id": item.author_actor_id,
                    "resolution": item.resolution,
                }
                for item in comments
            ],
            "decisions": [
                {
                    "id": item.id,
                    "sequence": item.sequence,
                    "decision": item.decision,
                    "rationale": item.rationale,
                    "decided_by_actor_id": item.decided_by_actor_id,
                    "decided_at": item.decided_at,
                    "is_final": item.is_final,
                }
                for item in decisions
            ],
        },
        "conditions_to_close": [
            {
                "id": item.id,
                "description": item.description,
                "owner_actor_id": item.owner_actor_id,
                "due_date": item.due_date,
                "status": item.status,
                "evidence_refs": item.evidence_refs,
            }
            for item in conditions
        ],
    }
    return packet, deal, _safe_json(payload)


def _render_json(payload: dict) -> bytes:
    return json.dumps(payload, indent=2, sort_keys=True, ensure_ascii=False).encode("utf-8")


def _render_xlsx(payload: dict) -> bytes:
    workbook = Workbook()
    workbook.remove(workbook.active)
    header_fill = PatternFill("solid", fgColor="17243D")
    for section, value in payload.items():
        title = re.sub(r"[^A-Za-z0-9 _-]", "", section.replace("_", " ").title())[:31]
        sheet = workbook.create_sheet(title or "Section")
        sheet.append(["Field", "Value"])
        for cell in sheet[1]:
            cell.font = Font(color="FFFFFF", bold=True)
            cell.fill = header_fill
        for field, item in _flatten(value):
            sheet.append([_excel_safe(field), _excel_safe(item)])
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = f"A1:B{max(1, sheet.max_row)}"
        sheet.column_dimensions["A"].width = 48
        sheet.column_dimensions["B"].width = 90
        for row in sheet.iter_rows(min_row=2, max_col=2):
            row[1].alignment = Alignment(wrap_text=True, vertical="top")
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _render_docx(payload: dict) -> bytes:
    document = Document()
    metadata = payload["export_metadata"]
    document.add_heading(f"Investment Committee Pack — {metadata['deal_name']}", 0)
    document.add_paragraph(
        f"Packet {metadata['packet_id']} · version {metadata['packet_version']} · "
        f"status {metadata['packet_status']}"
    )
    document.add_paragraph(f"Content hash: {metadata['packet_content_hash']}")
    for section, value in payload.items():
        document.add_heading(section.replace("_", " ").title(), level=1)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Field"
        table.rows[0].cells[1].text = "Value"
        for field, item in _flatten(value):
            cells = table.add_row().cells
            cells[0].text = field
            cells[1].text = item
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_pdf(payload: dict) -> bytes:
    buffer = io.BytesIO()
    metadata = payload["export_metadata"]
    styles = getSampleStyleSheet()
    story: list[Any] = [
        Paragraph(f"Investment Committee Pack — {escape(metadata['deal_name'])}", styles["Title"]),
        Spacer(1, 0.15 * inch),
    ]
    for section, value in payload.items():
        story.append(Paragraph(escape(section.replace("_", " ").title()), styles["Heading2"]))
        rows = [[Paragraph("Field", styles["BodyText"]), Paragraph("Value", styles["BodyText"])]]
        for field, item in _flatten(value):
            rows.append([
                Paragraph(escape(field), styles["BodyText"]),
                Paragraph(escape(item[:2_000]), styles["BodyText"]),
            ])
        table = Table(rows, colWidths=[2.25 * inch, 4.75 * inch], repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#17243D")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.extend([table, Spacer(1, 0.18 * inch)])

    def footer(canvas, document):
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        label = (
            f"Packet {metadata['packet_id']} · v{metadata['packet_version']} · "
            f"{metadata['packet_content_hash']} · page {document.page}"
        )
        canvas.drawString(0.55 * inch, 0.35 * inch, label[:150])
        canvas.restoreState()

    document = SimpleDocTemplate(
        buffer, pagesize=LETTER, leftMargin=0.5 * inch, rightMargin=0.5 * inch,
        topMargin=0.5 * inch, bottomMargin=0.55 * inch,
    )
    document.build(story, onFirstPage=footer, onLaterPages=footer)
    return buffer.getvalue()


def render_and_record_export(
    session: Session,
    packet_id: str,
    request: ExportRequest,
    actor: ActorContext | None = None,
) -> ExportedICFile:
    packet, deal, payload = _packet_payload(session, packet_id, actor)
    renderer = {
        "json": _render_json,
        "xlsx": _render_xlsx,
        "docx": _render_docx,
        "pdf": _render_pdf,
    }[request.format]
    content = renderer(payload)
    digest = hashlib.sha256(content).hexdigest()
    filename = f"{_file_stem(deal.code or deal.name)}-ic-v{packet.version}.{request.format}"
    manifest = {
        "schema_version": "1.0",
        "format": request.format,
        "filename": filename,
        "media_type": _MEDIA_TYPES[request.format],
        "byte_size": len(content),
        "file_sha256": digest,
        "packet_id": packet.id,
        "packet_version": packet.version,
        "packet_content_hash": packet.content_hash,
        "generated_at": payload["export_metadata"]["generated_at"],
        "sections": list(payload),
    }
    # G74: manifest_hash is the SHA-256 of the canonical manifest bytes EXCLUDING the attestation
    # block (identical to the pre-attestation canonicalization), and the Ed25519 signature —
    # when a key is configured — covers those exact same bytes, so hash and signature attest the
    # same thing and the signature is never self-referential. No key => honest "unsigned" block.
    core_bytes = export_signing_service.canonical_manifest_bytes(manifest)
    manifest["attestation"] = export_signing_service.build_attestation(
        core_bytes,
        signed_payload=(
            "canonical manifest JSON excluding the attestation block "
            "(sort_keys=True, separators=(',', ':')); its SHA-256 is manifest_hash"
        ),
    )
    record = ICPacketExport(
        packet_id=packet.id,
        format=request.format,
        manifest=manifest,
        manifest_hash=hashlib.sha256(core_bytes).hexdigest(),
        requested_by_actor_id=actor.actor_id if actor else None,
    )
    session.add(record)
    session.flush()
    workflow._audit(
        session,
        deal.organization_id,
        deal.id,
        actor,
        "ic_packet.export_file_created",
        record,
        {"format": request.format, "file_sha256": digest, "byte_size": len(content)},
    )
    workflow._commit(session, record)
    return ExportedICFile(
        content=content,
        filename=filename,
        media_type=_MEDIA_TYPES[request.format],
        sha256=digest,
        export_id=record.id,
    )
