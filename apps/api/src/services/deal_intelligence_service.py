"""Offline-testable deal-room parsing, cited retrieval, extraction, review, and diffing."""
from __future__ import annotations

import csv
import hashlib
import io
import json
import re
import unicodedata
import zipfile
from copy import deepcopy
from datetime import datetime, timezone
from typing import Any, Iterable

from docx import Document as WordDocument
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session

from src.db.base import new_uuid
from src.models.deal_intelligence import (
    CitedQARun,
    ClaimReviewEvent,
    DataRoomChunk,
    DataRoomDocument,
    DocumentComparison,
    IntelligenceEvaluation,
    SecFilingComparison,
    StructuredClaim,
)
from src.models.document import DocumentChunk as FilingChunk
from src.models.deal_workflow import Deal
from src.models.evidence import Evidence
from src.models.filing import Filing
from src.models.workspace import Workspace
from src.schemas.deal_intelligence import (
    CitedQARequest,
    ClaimReviewRequest,
    ComparisonRequest,
    DocumentTextCreate,
    EvaluationRequest,
    ExtractionRequest,
    QAFilters,
    SecFilingComparisonRequest,
)
from src.schemas.deal_workflow import ActorContext
from src.services.common import insert_versioned

MAX_DOCUMENT_BYTES = 20 * 1024 * 1024
MAX_CHUNK_CHARS = 4_000
QA_ALGORITHM_VERSION = "lexical-cited-v2"
EXTRACTION_VERSION = "rules-evidence-v1"
COMPARISON_VERSION = "locator-diff-v1"
EVALUATION_VERSION = "evidence-eval-v1"
ABSTENTION = "I could not find a supported answer in the selected documents."

CONTENT_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".csv": "text/csv",
    ".txt": "text/plain",
}

_STOPWORDS = {
    "a",
    "an",
    "and",
    "are",
    "as",
    "at",
    "be",
    "by",
    "did",
    "do",
    "does",
    "for",
    "from",
    "had",
    "has",
    "have",
    "how",
    "in",
    "is",
    "it",
    "of",
    "on",
    "or",
    "the",
    "their",
    "to",
    "was",
    "were",
    "what",
    "when",
    "where",
    "which",
    "who",
    "with",
}

_CATEGORY_RULES: dict[str, dict[str, tuple[str, ...]]] = {
    "debt_term": {
        "covenant_term": ("covenant", "headroom", "fixed charge coverage"),
        "interest_rate": ("interest rate", "sofr", "spread", "coupon", "basis points"),
        "maturity": ("maturity", "matures", "due date"),
        "leverage": ("leverage", "debt to ebitda"),
        "debt_facility": ("debt", "loan", "facility", "revolver", "term loan", "borrowings"),
    },
    "customer": {
        "customer_concentration": (
            "customer concentration",
            "top customer",
            "largest customer",
            "top 10 customers",
        ),
        "customer_retention": ("retention", "churn", "renewal rate", "net revenue retention"),
        "customer_count": ("customers", "customer count", "accounts"),
    },
    "contract": {
        "change_of_control": ("change of control", "assignment consent"),
        "termination_term": ("termination", "terminate", "notice period"),
        "renewal_term": ("renewal", "auto-renew", "expiration", "expires"),
        "contract_term": ("contract", "agreement", "term of"),
    },
    "kpi": {
        "arr": ("arr", "annual recurring revenue"),
        "mrr": ("mrr", "monthly recurring revenue"),
        "revenue": ("revenue", "sales"),
        "ebitda": ("ebitda",),
        "gross_margin": ("gross margin",),
        "retention": ("retention", "nrr", "grr"),
        "growth": ("growth", "grew", "increase", "decrease"),
    },
    "qoe_candidate": {
        "add_back": ("add-back", "add back", "adjusted ebitda"),
        "non_recurring": ("non-recurring", "nonrecurring", "one-time", "one time"),
        "run_rate": ("run-rate", "run rate", "pro forma"),
        "owner_compensation": ("owner compensation", "management fee", "related party"),
        "qoe_adjustment": ("quality of earnings", "qoe", "normalization", "adjustment"),
    },
}

_NUMBER_RE = re.compile(
    r"(?P<currency>\$|USD\s*|EUR\s*|GBP\s*)?"
    r"(?P<number>-?\d[\d,]*(?:\.\d+)?)\s*"
    r"(?P<scale>billion|million|thousand|bn|mm|m|k)?\s*"
    r"(?P<unit>%|x|bps|basis points)?",
    re.IGNORECASE,
)
_PERIOD_RE = re.compile(
    r"\b(?:FY\s*20\d{2}|Q[1-4]\s*20\d{2}|LTM|TTM|20\d{2}(?:\s*[-/]\s*20\d{2})?)\b",
    re.IGNORECASE,
)


class IntelligenceError(ValueError):
    def __init__(self, message: str, *, status_code: int = 400) -> None:
        super().__init__(message)
        self.message = message
        self.status_code = status_code


class IntelligenceNotFound(IntelligenceError):
    def __init__(self, message: str) -> None:
        super().__init__(message, status_code=404)


class IntelligenceConflict(IntelligenceError):
    def __init__(self, message: str) -> None:
        super().__init__(message, status_code=409)


def _actor_id(actor: ActorContext | None) -> str | None:
    return actor.actor_id if actor else None


def _require_actor(actor: ActorContext | None, action: str) -> str:
    actor_id = _actor_id(actor)
    if not actor_id:
        raise IntelligenceError(f"{action} requires an authenticated actor", status_code=401)
    return actor_id


def _require_human_reviewer(actor: ActorContext | None, action: str) -> str:
    """Four-eyes reviews demand a server-verified human identity: the trusted-service path's
    actor id is caller-chosen, so one automation token could satisfy both sides of the rule."""
    actor_id = _require_actor(actor, action)
    if actor is not None and actor.via_trusted_service:
        raise IntelligenceError(
            f"{action} requires a human user session; trusted-service automation "
            "cannot act as a reviewer",
            status_code=403,
        )
    return actor_id


def _verify_scope(actor: ActorContext | None, organization_id: str) -> None:
    if actor and actor.organization_id and actor.organization_id != organization_id:
        raise IntelligenceError(
            "Organization scope does not permit this operation", status_code=403
        )


def _deal(session: Session, deal_id: str, actor: ActorContext | None = None) -> Deal:
    deal = session.get(Deal, deal_id)
    # Cross-tenant lookups return the SAME 404 as an unknown id (mirroring ``_document``): a 403
    # here would confirm the deal id exists in another tenant.
    if deal is None or (
        actor and actor.organization_id and actor.organization_id != deal.organization_id
    ):
        raise IntelligenceNotFound(f"Deal '{deal_id}' not found")
    return deal


def _document(
    session: Session, document_id: str, actor: ActorContext | None = None
) -> DataRoomDocument:
    # Scope the lookup in SQL when a verified organization is available.  Besides avoiding a
    # cross-tenant raw-bytes load, returning the same 404 as an unknown id prevents callers from
    # using get/chunk/download responses as a document-id oracle.
    statement = select(DataRoomDocument).where(DataRoomDocument.id == document_id)
    if actor and actor.organization_id:
        statement = statement.join(Deal, Deal.id == DataRoomDocument.deal_id).where(
            Deal.organization_id == actor.organization_id
        )
    document = session.scalar(statement)
    if document is None:
        raise IntelligenceNotFound("Document not found")
    _deal(session, document.deal_id, actor)
    return document


def _claim(session: Session, claim_id: str, actor: ActorContext | None = None) -> StructuredClaim:
    claim = session.get(StructuredClaim, claim_id)
    if claim is None:
        raise IntelligenceNotFound(f"Claim '{claim_id}' not found")
    _deal(session, claim.deal_id, actor)
    return claim


def _commit(session: Session, entity: Any) -> Any:
    try:
        session.commit()
    except IntegrityError as exc:
        session.rollback()
        raise IntelligenceConflict("The immutable intelligence record conflicts with existing data") from exc
    return entity


def _sha256_bytes(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def _sha256_text(content: str) -> str:
    return hashlib.sha256(content.encode("utf-8")).hexdigest()


def _utc_iso(value: datetime) -> str:
    if value.tzinfo is None:
        value = value.replace(tzinfo=timezone.utc)
    return value.astimezone(timezone.utc).isoformat()


def _normalized(text: str) -> str:
    return " ".join(unicodedata.normalize("NFKC", text).casefold().split())


def secure_filename(filename: str) -> tuple[str, str]:
    """Return a display-safe basename and validated extension; never return a path."""
    normalized = unicodedata.normalize("NFKC", filename).replace("\\", "/")
    basename = normalized.rsplit("/", 1)[-1].strip().strip(".")
    basename = re.sub(r"[^A-Za-z0-9._() -]+", "_", basename)
    basename = re.sub(r"\s+", " ", basename).strip(" .")
    if not basename:
        raise IntelligenceError("A valid filename is required", status_code=422)
    if len(basename) > 240:
        stem, dot, suffix = basename.rpartition(".")
        basename = f"{stem[:220]}{dot}{suffix}" if dot else basename[:240]
    extension = "." + basename.rsplit(".", 1)[-1].lower() if "." in basename else ""
    if extension not in CONTENT_TYPES:
        supported = ", ".join(sorted(CONTENT_TYPES))
        raise IntelligenceError(
            f"Unsupported document type '{extension or '(none)'}'; expected {supported}",
            status_code=415,
        )
    return basename, extension


def _validate_content(
    filename: str, content: bytes, content_type: str | None
) -> tuple[str, str, str]:
    safe_name, extension = secure_filename(filename)
    if not content:
        raise IntelligenceError("Document is empty", status_code=422)
    if len(content) > MAX_DOCUMENT_BYTES:
        raise IntelligenceError(
            f"Document exceeds the {MAX_DOCUMENT_BYTES // (1024 * 1024)} MB limit",
            status_code=413,
        )
    canonical_type = CONTENT_TYPES[extension]
    supplied_type = (content_type or "").split(";", 1)[0].strip().lower()
    permitted = {canonical_type, "application/octet-stream"}
    if extension in {".txt", ".csv"}:
        permitted.update({"text/plain", "text/csv", "application/csv", "application/vnd.ms-excel"})
    if supplied_type and supplied_type not in permitted:
        raise IntelligenceError(
            f"Content type '{supplied_type}' does not match {extension}", status_code=415
        )
    if extension == ".pdf" and not content.lstrip().startswith(b"%PDF-"):
        raise IntelligenceError("File content is not a PDF", status_code=422)
    if extension in {".txt", ".csv"} and b"\x00" in content:
        raise IntelligenceError("Text documents may not contain NUL bytes", status_code=422)
    if extension in {".docx", ".xlsx"}:
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                entries = archive.infolist()
                names = {entry.filename for entry in entries}
                expanded_size = sum(entry.file_size for entry in entries)
        except zipfile.BadZipFile as exc:
            raise IntelligenceError(f"File content is not a valid {extension}", status_code=422) from exc
        required_member = "word/document.xml" if extension == ".docx" else "xl/workbook.xml"
        if required_member not in names:
            raise IntelligenceError(f"File content does not match {extension}", status_code=422)
        if len(entries) > 10_000 or expanded_size > 100 * 1024 * 1024:
            raise IntelligenceError("Compressed document expands beyond safety limits", status_code=413)
    return safe_name, extension, canonical_type


def _segments(text: str, locator: dict[str, Any]) -> Iterable[tuple[str, dict[str, Any]]]:
    cleaned = text.replace("\x00", "").strip()
    if not cleaned:
        return
    if len(cleaned) <= MAX_CHUNK_CHARS:
        yield cleaned, locator
        return
    start = 0
    while start < len(cleaned):
        end = min(start + MAX_CHUNK_CHARS, len(cleaned))
        if end < len(cleaned):
            split = cleaned.rfind(" ", start, end)
            if split > start + MAX_CHUNK_CHARS // 2:
                end = split
        child_locator = {**locator, "char_start": start, "char_end": end}
        yield cleaned[start:end].strip(), child_locator
        start = end


def _parse_text(content: bytes) -> list[tuple[str, str, dict[str, Any]]]:
    try:
        text = content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise IntelligenceError("Text documents must be UTF-8 encoded", status_code=422) from exc
    units = [part for part in re.split(r"\n\s*\n|\r?\n", text) if part.strip()]
    output: list[tuple[str, str, dict[str, Any]]] = []
    for paragraph, unit in enumerate(units, 1):
        locator = {"type": "text", "paragraph": paragraph}
        output.extend(("text", value, place) for value, place in _segments(unit, locator))
    return output


def _row_text(values: list[Any], row_number: int) -> tuple[str, str, str] | None:
    populated = [(index, value) for index, value in enumerate(values, 1) if value not in (None, "")]
    if not populated:
        return None
    start_column = populated[0][0]
    end_column = populated[-1][0]
    cells = [f"{get_column_letter(index)}{row_number}={value}" for index, value in populated]
    return (
        "; ".join(cells),
        f"{get_column_letter(start_column)}{row_number}",
        f"{get_column_letter(end_column)}{row_number}",
    )


def _parse_csv(content: bytes) -> list[tuple[str, str, dict[str, Any]]]:
    try:
        text = content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise IntelligenceError("CSV documents must be UTF-8 encoded", status_code=422) from exc
    output: list[tuple[str, str, dict[str, Any]]] = []
    try:
        rows = csv.reader(io.StringIO(text))
        for row_number, row in enumerate(rows, 1):
            parsed = _row_text(row, row_number)
            if parsed is None:
                continue
            value, start, end = parsed
            locator = {
                "type": "csv",
                "sheet": "CSV",
                "row": row_number,
                "cell_start": start,
                "cell_end": end,
                "cell_range": f"{start}:{end}",
            }
            output.extend(("csv", item, place) for item, place in _segments(value, locator))
    except csv.Error as exc:
        raise IntelligenceError(f"Invalid CSV document: {exc}", status_code=422) from exc
    return output


def _parse_pdf(content: bytes) -> list[tuple[str, str, dict[str, Any]]]:
    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:
        raise IntelligenceError("Invalid or encrypted PDF document", status_code=422) from exc
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:
            raise IntelligenceError("Encrypted PDF documents are not supported", status_code=422) from exc
    output: list[tuple[str, str, dict[str, Any]]] = []
    for page_number, page in enumerate(reader.pages, 1):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:
            raise IntelligenceError(
                f"Could not extract PDF page {page_number}", status_code=422
            ) from exc
        paragraphs = [value for value in re.split(r"\n\s*\n|\r?\n", page_text) if value.strip()]
        for paragraph, value in enumerate(paragraphs, 1):
            locator = {"type": "pdf", "page": page_number, "paragraph": paragraph}
            output.extend(("pdf", item, place) for item, place in _segments(value, locator))
    return output


def _parse_docx(content: bytes) -> list[tuple[str, str, dict[str, Any]]]:
    try:
        document = WordDocument(io.BytesIO(content))
    except Exception as exc:
        raise IntelligenceError("Invalid DOCX document", status_code=422) from exc
    output: list[tuple[str, str, dict[str, Any]]] = []
    for paragraph, node in enumerate(document.paragraphs, 1):
        locator = {"type": "docx", "paragraph": paragraph}
        output.extend(("docx", item, place) for item, place in _segments(node.text, locator))
    for table_number, table in enumerate(document.tables, 1):
        for row_number, row in enumerate(table.rows, 1):
            values = [cell.text.strip() for cell in row.cells]
            parsed = _row_text(values, row_number)
            if parsed is None:
                continue
            value, start, end = parsed
            locator = {
                "type": "docx_table",
                "table": table_number,
                "row": row_number,
                "cell_start": start,
                "cell_end": end,
                "cell_range": f"{start}:{end}",
            }
            output.extend(("docx_table", item, place) for item, place in _segments(value, locator))
    return output


def _parse_xlsx(content: bytes) -> list[tuple[str, str, dict[str, Any]]]:
    try:
        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    except Exception as exc:
        raise IntelligenceError("Invalid XLSX document", status_code=422) from exc
    output: list[tuple[str, str, dict[str, Any]]] = []
    try:
        for sheet in workbook.worksheets:
            for row_number, row in enumerate(sheet.iter_rows(values_only=True), 1):
                parsed = _row_text(list(row), row_number)
                if parsed is None:
                    continue
                value, start, end = parsed
                locator = {
                    "type": "xlsx",
                    "sheet": sheet.title,
                    "row": row_number,
                    "cell_start": start,
                    "cell_end": end,
                    "cell_range": f"{start}:{end}",
                }
                output.extend(("xlsx", item, place) for item, place in _segments(value, locator))
    finally:
        workbook.close()
    return output


def extract_chunks(content: bytes, extension: str) -> list[tuple[str, str, dict[str, Any]]]:
    parser = {
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
        ".xlsx": _parse_xlsx,
        ".csv": _parse_csv,
        ".txt": _parse_text,
    }[extension]
    return parser(content)


def ingest_document(
    session: Session,
    deal_id: str,
    *,
    filename: str,
    content: bytes,
    content_type: str | None = None,
    title: str | None = None,
    logical_document_id: str | None = None,
    document_metadata: dict[str, Any] | None = None,
    actor: ActorContext | None = None,
) -> DataRoomDocument:
    _deal(session, deal_id, actor)
    safe_name, extension, canonical_type = _validate_content(filename, content, content_type)
    if logical_document_id and not re.fullmatch(r"[A-Za-z0-9_-]{1,32}", logical_document_id):
        raise IntelligenceError(
            "logical_document_id may contain only letters, numbers, underscores, and hyphens",
            status_code=422,
        )
    try:
        json.dumps(document_metadata or {})
    except (TypeError, ValueError) as exc:
        raise IntelligenceError("document_metadata must be JSON serializable", status_code=422) from exc

    parsed_chunks = extract_chunks(content, extension)
    logical_id = logical_document_id or new_uuid()
    digest = _sha256_bytes(content)

    def _build_document() -> DataRoomDocument:
        previous = session.scalar(
            select(DataRoomDocument)
            .where(
                DataRoomDocument.deal_id == deal_id,
                DataRoomDocument.logical_document_id == logical_id,
            )
            .order_by(DataRoomDocument.version.desc())
        )
        # Inside the retry loop deliberately: a concurrent upload that wins the version race may
        # have identical content, and the retry must then surface the same 409 as a serial upload.
        if previous and previous.sha256 == digest:
            raise IntelligenceConflict(
                f"Document content is identical to logical document {logical_id} version {previous.version}"
            )
        return DataRoomDocument(
            deal_id=deal_id,
            logical_document_id=logical_id,
            version=(previous.version + 1) if previous else 1,
            supersedes_document_id=previous.id if previous else None,
            title=(title or safe_name).strip()[:240],
            filename=safe_name,
            original_filename=filename[:500],
            extension=extension,
            content_type=canonical_type,
            sha256=digest,
            byte_size=len(content),
            raw_bytes=content,
            document_metadata=deepcopy(document_metadata or {}),
            uploaded_by_actor_id=_actor_id(actor),
        )

    document = insert_versioned(session, _build_document)
    for ordinal, (locator_type, text, locator) in enumerate(parsed_chunks, 1):
        normalized = _normalized(text)
        session.add(
            DataRoomChunk(
                deal_id=deal_id,
                document_id=document.id,
                ordinal=ordinal,
                locator_type=locator_type,
                locator=locator,
                text=text,
                normalized_text=normalized,
                content_hash=_sha256_text(text),
                char_count=len(text),
            )
        )
    return _commit(session, document)


def ingest_text_document(
    session: Session,
    deal_id: str,
    data: DocumentTextCreate,
    actor: ActorContext | None = None,
) -> DataRoomDocument:
    _, extension = secure_filename(data.filename)
    if extension not in {".txt", ".csv"}:
        raise IntelligenceError(
            "JSON text ingestion supports .txt and .csv; use multipart upload for binary files",
            status_code=415,
        )
    return ingest_document(
        session,
        deal_id,
        filename=data.filename,
        content=data.text.encode("utf-8"),
        content_type=data.content_type,
        title=data.title,
        logical_document_id=data.logical_document_id,
        document_metadata=data.document_metadata,
        actor=actor,
    )


def list_documents(
    session: Session,
    deal_id: str,
    actor: ActorContext | None = None,
    *,
    latest_only: bool = True,
    logical_document_id: str | None = None,
) -> list[DataRoomDocument]:
    _deal(session, deal_id, actor)
    statement = select(DataRoomDocument).where(DataRoomDocument.deal_id == deal_id)
    if logical_document_id:
        statement = statement.where(DataRoomDocument.logical_document_id == logical_document_id)
    documents = list(
        session.scalars(
            statement.order_by(
                DataRoomDocument.logical_document_id, DataRoomDocument.version.desc()
            )
        )
    )
    if not latest_only:
        return documents
    latest: dict[str, DataRoomDocument] = {}
    for document in documents:
        latest.setdefault(document.logical_document_id, document)
    return list(latest.values())


def list_document_versions(
    session: Session,
    deal_id: str,
    logical_document_id: str,
    actor: ActorContext | None = None,
) -> list[DataRoomDocument]:
    documents = list_documents(
        session,
        deal_id,
        actor,
        latest_only=False,
        logical_document_id=logical_document_id,
    )
    if not documents:
        raise IntelligenceNotFound(f"Logical document '{logical_document_id}' not found")
    return sorted(documents, key=lambda item: item.version)


def get_document(
    session: Session, document_id: str, actor: ActorContext | None = None
) -> DataRoomDocument:
    return _document(session, document_id, actor)


def list_chunks(
    session: Session, document_id: str, actor: ActorContext | None = None
) -> list[DataRoomChunk]:
    document = _document(session, document_id, actor)
    return list(
        session.scalars(
            select(DataRoomChunk)
            .where(DataRoomChunk.document_id == document.id)
            .order_by(DataRoomChunk.ordinal)
        )
    )


def _metadata_matches(document: DataRoomDocument, filters: dict[str, Any]) -> bool:
    return all(document.document_metadata.get(key) == value for key, value in filters.items())


def _filtered_documents(
    session: Session,
    deal_id: str,
    filters: QAFilters,
    actor: ActorContext | None,
) -> list[DataRoomDocument]:
    documents = list_documents(session, deal_id, actor, latest_only=False)
    if filters.document_ids:
        allowed = set(filters.document_ids)
        documents = [item for item in documents if item.id in allowed]
    if filters.logical_document_ids:
        allowed = set(filters.logical_document_ids)
        documents = [item for item in documents if item.logical_document_id in allowed]
    if filters.extensions:
        allowed = {value if value.startswith(".") else f".{value}" for value in filters.extensions}
        documents = [item for item in documents if item.extension in allowed]
    if filters.versions:
        allowed_versions = set(filters.versions)
        documents = [item for item in documents if item.version in allowed_versions]
    if filters.metadata:
        documents = [item for item in documents if _metadata_matches(item, filters.metadata)]
    if filters.latest_only and not filters.document_ids and not filters.versions:
        latest: dict[str, DataRoomDocument] = {}
        for document in sorted(documents, key=lambda item: item.version, reverse=True):
            latest.setdefault(document.logical_document_id, document)
        documents = list(latest.values())
    return documents


def _tokens(text: str) -> set[str]:
    return {
        token
        for token in re.findall(r"[A-Za-z0-9]+", text.casefold())
        if len(token) > 1 and token not in _STOPWORDS
    }


def _sentence_spans(text: str) -> list[tuple[int, int, str]]:
    spans: list[tuple[int, int, str]] = []
    start_at = 0
    # Decimal points are not sentence boundaries: `$185.0 million` must remain an exact quote.
    boundaries = list(re.finditer(r"[.!?]+(?=\s+|$)|\n+", text))
    for boundary in boundaries:
        raw = text[start_at:boundary.end()]
        leading = len(raw) - len(raw.lstrip())
        trailing = len(raw) - len(raw.rstrip())
        start = start_at + leading
        end = boundary.end() - trailing
        if start < end:
            spans.append((start, end, text[start:end]))
        start_at = boundary.end()
    if start_at < len(text):
        raw = text[start_at:]
        leading = len(raw) - len(raw.lstrip())
        trailing = len(raw) - len(raw.rstrip())
        start = start_at + leading
        end = len(text) - trailing
        if start < end:
            spans.append((start, end, text[start:end]))
    return spans or [(0, len(text), text)]


def _citation(
    document: DataRoomDocument, chunk: DataRoomChunk, quote: str
) -> dict[str, Any]:
    return {
        "document_id": document.id,
        "logical_document_id": document.logical_document_id,
        "document_version": document.version,
        "filename": document.filename,
        "sha256": document.sha256,
        "chunk_id": chunk.id,
        "content_hash": chunk.content_hash,
        "locator": deepcopy(chunk.locator),
        "quote": quote,
    }


def answer_question(
    session: Session,
    deal_id: str,
    data: CitedQARequest,
    actor: ActorContext | None = None,
) -> CitedQARun:
    documents = _filtered_documents(session, deal_id, data.filters, actor)
    question_tokens = _tokens(data.question)
    candidates: list[tuple[float, DataRoomDocument, DataRoomChunk, str, set[str]]] = []
    chunk_count = 0
    for document in documents:
        chunks = list(
            session.scalars(
                select(DataRoomChunk)
                .where(DataRoomChunk.document_id == document.id)
                .order_by(DataRoomChunk.ordinal)
            )
        )
        chunk_count += len(chunks)
        for chunk in chunks:
            for _, _, sentence in _sentence_spans(chunk.text):
                sentence_tokens = _tokens(sentence)
                overlap = question_tokens & sentence_tokens
                if not overlap:
                    continue
                coverage = len(overlap) / max(len(question_tokens), 1)
                density = len(overlap) / max(len(sentence_tokens), 1)
                score = len(overlap) * 2 + coverage + density
                candidates.append((score, document, chunk, sentence[:800], overlap))

    if not candidates:
        status = "abstained"
        answer = ABSTENTION
        citations: list[dict[str, Any]] = []
        retrieval = {
            "candidate_document_count": len(documents),
            "candidate_chunk_count": chunk_count,
            "matched_terms": [],
            "score": 0.0,
            "abstention_reason": "no lexical evidence matched the question",
        }
    else:
        # Greedily cover compound questions with up to three independently cited sentences.
        # Each selected sentence must add a question term that earlier sentences did not cover.
        remaining = set(question_tokens)
        selected: list[tuple[float, DataRoomDocument, DataRoomChunk, str, set[str]]] = []
        pool = sorted(
            candidates,
            key=lambda item: (item[0], -item[2].ordinal, item[1].version),
            reverse=True,
        )
        while pool and len(selected) < 3:
            best = max(
                pool,
                key=lambda item: (
                    len(item[4] & remaining), item[0], -item[2].ordinal, item[1].version
                ),
            )
            if not (best[4] & remaining):
                break
            selected.append(best)
            remaining -= best[4]
            pool.remove(best)
        score = sum(item[0] for item in selected)
        overlap = set().union(*(item[4] for item in selected))
        answer = " ".join(item[3].strip() for item in selected)
        status = "answered"
        citations = [_citation(item[1], item[2], item[3]) for item in selected]
        retrieval = {
            "candidate_document_count": len(documents),
            "candidate_chunk_count": chunk_count,
            "matched_terms": sorted(overlap),
            "score": round(score, 6),
            "abstention_reason": None,
        }
    run = CitedQARun(
        deal_id=deal_id,
        question=data.question,
        filters=data.filters.model_dump(mode="json"),
        status=status,
        answer=answer,
        citations=citations,
        retrieval_metadata=retrieval,
        answer_hash=_sha256_text(answer),
        algorithm_version=QA_ALGORITHM_VERSION,
        created_by_actor_id=_actor_id(actor),
    )
    session.add(run)
    return _commit(session, run)


def list_qa_runs(
    session: Session,
    deal_id: str,
    actor: ActorContext | None = None,
    *,
    limit: int = 100,
) -> list[CitedQARun]:
    _deal(session, deal_id, actor)
    return list(
        session.scalars(
            select(CitedQARun)
            .where(CitedQARun.deal_id == deal_id)
            .order_by(CitedQARun.created_at.desc())
            .limit(min(max(limit, 1), 1_000))
        )
    )


def _numeric_attributes(sentence: str) -> tuple[float | None, str | None, str | None]:
    matches = list(_NUMBER_RE.finditer(sentence))
    if not matches:
        return None, None, None

    def priority(match: re.Match[str]) -> tuple[int, int]:
        groups = match.groupdict()
        explicit = int(bool(groups["currency"] or groups["scale"] or groups["unit"]))
        return explicit, -match.start()

    selected = max(matches, key=priority)
    groups = selected.groupdict()
    try:
        number = float(groups["number"].replace(",", ""))
    except ValueError:
        number = None
    scale = (groups["scale"] or "").lower()
    raw_unit = (groups["unit"] or "").lower()
    currency_token = (groups["currency"] or "").strip().upper()
    currency = {"$": "USD", "USD": "USD", "EUR": "EUR", "GBP": "GBP"}.get(currency_token)
    scale_unit = {
        "billion": "billions",
        "bn": "billions",
        "million": "millions",
        "mm": "millions",
        "m": "millions",
        "thousand": "thousands",
        "k": "thousands",
    }.get(scale)
    if raw_unit == "%":
        unit = "percent"
    elif raw_unit == "x":
        unit = "turns"
    elif raw_unit in {"bps", "basis points"}:
        unit = "basis_points"
    elif currency and scale_unit:
        unit = f"{currency}_{scale_unit}"
    elif currency:
        unit = currency
    else:
        unit = scale_unit
    return number, unit, currency


def _period(sentence: str) -> str | None:
    match = _PERIOD_RE.search(sentence)
    return re.sub(r"\s+", "", match.group(0).upper()) if match else None


def _field_for(category: str, normalized_sentence: str) -> tuple[str, int] | None:
    for field_name, keywords in _CATEGORY_RULES[category].items():
        matched = sum(1 for keyword in keywords if keyword in normalized_sentence)
        if matched:
            return field_name, matched
    return None


def _documents_for_extraction(
    session: Session,
    deal_id: str,
    request: ExtractionRequest,
    actor: ActorContext | None,
) -> list[DataRoomDocument]:
    filters = QAFilters(
        document_ids=request.document_ids,
        latest_only=request.latest_only,
    )
    return _filtered_documents(session, deal_id, filters, actor)


def extract_structured_claims(
    session: Session,
    deal_id: str,
    data: ExtractionRequest,
    actor: ActorContext | None = None,
) -> list[StructuredClaim]:
    documents = _documents_for_extraction(session, deal_id, data, actor)
    existing = list(
        session.scalars(select(StructuredClaim).where(StructuredClaim.deal_id == deal_id))
    )
    signatures: dict[tuple[Any, ...], StructuredClaim] = {}
    for claim in existing:
        if claim.revision == 1:
            signatures[
                (
                    claim.chunk_id,
                    claim.category,
                    claim.field_name,
                    claim.source_span.get("start"),
                    claim.source_span.get("end"),
                )
            ] = claim

    results: list[StructuredClaim] = []
    for document in documents:
        chunks = list(
            session.scalars(
                select(DataRoomChunk)
                .where(DataRoomChunk.document_id == document.id)
                .order_by(DataRoomChunk.ordinal)
            )
        )
        for chunk in chunks:
            for start, end, sentence in _sentence_spans(chunk.text):
                normalized_sentence = _normalized(sentence)
                for category in data.categories:
                    field_match = _field_for(category, normalized_sentence)
                    if field_match is None:
                        continue
                    field_name, keyword_count = field_match
                    value_number, unit, currency = _numeric_attributes(sentence)
                    confidence = min(
                        0.96,
                        0.7 + min(keyword_count, 2) * 0.06 + (0.08 if value_number is not None else 0),
                    )
                    if confidence < data.min_confidence:
                        continue
                    signature = (chunk.id, category, field_name, start, end)
                    if signature in signatures:
                        results.append(signatures[signature])
                        continue
                    claim = StructuredClaim(
                        deal_id=deal_id,
                        logical_claim_id=new_uuid(),
                        revision=1,
                        document_id=document.id,
                        chunk_id=chunk.id,
                        category=category,
                        field_name=field_name,
                        value_text=sentence,
                        value_number=value_number,
                        unit=unit,
                        period=_period(sentence),
                        currency=currency,
                        confidence=round(confidence, 3),
                        source_locator=deepcopy(chunk.locator),
                        source_span={"start": start, "end": end, "text": sentence},
                        review_status="unreviewed",
                        extraction_version=EXTRACTION_VERSION,
                        created_by_actor_id=_actor_id(actor),
                    )
                    session.add(claim)
                    session.flush()
                    signatures[signature] = claim
                    results.append(claim)
    session.commit()
    return results


def _latest_claim_revision(session: Session, logical_claim_id: str) -> StructuredClaim | None:
    return session.scalar(
        select(StructuredClaim)
        .where(StructuredClaim.logical_claim_id == logical_claim_id)
        .order_by(StructuredClaim.revision.desc())
    )


def _resolved_claim_source(
    session: Session, claim: StructuredClaim
) -> tuple[DataRoomDocument, DataRoomChunk, str]:
    """Resolve the exact immutable document/chunk/span that supports a claim."""
    document = session.get(DataRoomDocument, claim.document_id)
    chunk = session.get(DataRoomChunk, claim.chunk_id)
    if (
        document is None
        or document.deal_id != claim.deal_id
        or chunk is None
        or chunk.deal_id != claim.deal_id
        or chunk.document_id != document.id
    ):
        raise IntelligenceConflict(f"Claim '{claim.id}' source provenance no longer resolves")
    span = claim.source_span or {}
    quoted = span.get("text", "")
    start = span.get("start")
    end = span.get("end")
    if (
        not quoted
        or not isinstance(start, int)
        or isinstance(start, bool)
        or not isinstance(end, int)
        or isinstance(end, bool)
        or start < 0
        or end <= start
        or chunk.text[start:end] != quoted
    ):
        raise IntelligenceConflict(f"Claim '{claim.id}' source span no longer resolves")
    return document, chunk, quoted


def _promote_approved_claim(
    session: Session,
    deal: Deal,
    claim: StructuredClaim,
    approval: ClaimReviewEvent,
    document: DataRoomDocument,
    chunk: DataRoomChunk,
    quoted: str,
) -> Evidence | None:
    """Bridge an approved private claim into the workspace's governed Evidence ledger."""
    if not deal.workspace_id:
        return None
    source_url = f"dealroom://claims/{claim.id}"
    existing = session.scalar(
        select(Evidence).where(
            Evidence.workspace_id == deal.workspace_id,
            Evidence.source_type == "approved_private_claim",
            Evidence.source_url == source_url,
        )
    )
    if existing is not None:
        return existing

    # Evidence has a deliberately compact generic schema.  Preserve the machine-resolvable
    # private provenance in source_section while the IC manifest below repeats the binding as
    # structured JSON for downstream consumers.
    source_binding = {
        "claim_id": claim.id,
        "logical_claim_id": claim.logical_claim_id,
        "claim_revision": claim.revision,
        "review_event_id": approval.id,
        "document_id": document.id,
        "document_sha256": document.sha256,
        "chunk_id": chunk.id,
        "chunk_hash": chunk.content_hash,
        "locator": claim.source_locator,
        "span": claim.source_span,
    }
    from src.services import evidence_service

    return evidence_service.create(
        session,
        deal.workspace_id,
        claim=f"{claim.field_name}: {claim.value_text}",
        claim_type="fact",
        source_name=document.filename,
        source_type="approved_private_claim",
        source_url=source_url,
        source_date=approval.created_at.date().isoformat(),
        source_section=json.dumps(
            source_binding, sort_keys=True, separators=(",", ":"), default=str
        ),
        evidence_text=quoted,
        confidence=claim.confidence,
        agent_name="authenticated_claim_reviewer",
    )


def review_claim(
    session: Session,
    claim_id: str,
    data: ClaimReviewRequest,
    actor: ActorContext | None = None,
) -> tuple[StructuredClaim, ClaimReviewEvent]:
    source = _claim(session, claim_id, actor)
    reviewer_id = _require_human_reviewer(actor, "Claim review")
    latest = _latest_claim_revision(session, source.logical_claim_id)
    if latest is None or latest.id != source.id:
        latest_revision = latest.revision if latest else "unknown"
        raise IntelligenceConflict(
            f"Claim revision is stale; latest revision is {latest_revision}"
        )
    if data.expected_revision != source.revision:
        raise IntelligenceConflict(
            f"Claim revision changed from expected {data.expected_revision} to {source.revision}"
        )

    values = {
        "field_name": source.field_name,
        "value_text": source.value_text,
        "value_number": source.value_number,
        "unit": source.unit,
        "period": source.period,
        "currency": source.currency,
        "confidence": source.confidence,
    }
    changes: dict[str, Any] = {}
    resulting_status = {"approve": "approved", "reject": "rejected", "edit": "unreviewed"}[
        data.action
    ]
    if (
        data.action in {"approve", "reject"}
        and source.created_by_actor_id
        and source.created_by_actor_id == reviewer_id
    ):
        raise IntelligenceConflict("A distinct reviewer must approve or reject the claim")
    if data.action == "edit":
        for field_name in values:
            replacement = getattr(data, field_name)
            if replacement is not None and replacement != values[field_name]:
                changes[field_name] = {"from": values[field_name], "to": replacement}
                values[field_name] = replacement
    else:
        changes["review_status"] = {"from": source.review_status, "to": resulting_status}

    revision = StructuredClaim(
        deal_id=source.deal_id,
        logical_claim_id=source.logical_claim_id,
        revision=source.revision + 1,
        supersedes_claim_id=source.id,
        document_id=source.document_id,
        chunk_id=source.chunk_id,
        category=source.category,
        field_name=values["field_name"],
        value_text=values["value_text"],
        value_number=values["value_number"],
        unit=values["unit"],
        period=values["period"],
        currency=values["currency"],
        confidence=values["confidence"],
        source_locator=deepcopy(source.source_locator),
        source_span=deepcopy(source.source_span),
        review_status=resulting_status,
        extraction_version=source.extraction_version,
        created_by_actor_id=reviewer_id,
    )
    session.add(revision)
    session.flush()
    review = ClaimReviewEvent(
        deal_id=source.deal_id,
        logical_claim_id=source.logical_claim_id,
        from_claim_id=source.id,
        to_claim_id=revision.id,
        from_revision=source.revision,
        to_revision=revision.revision,
        action=data.action,
        prior_status=source.review_status,
        resulting_status=resulting_status,
        changes=changes,
        note=data.note,
        reviewer_actor_id=reviewer_id,
    )
    session.add(review)
    session.flush()
    if data.action == "approve":
        deal = _deal(session, source.deal_id, actor)
        document, chunk, quoted = _resolved_claim_source(session, revision)
        _promote_approved_claim(
            session, deal, revision, review, document, chunk, quoted
        )
    session.commit()
    return revision, review


def approved_claim_manifest(
    session: Session,
    deal_id: str,
    claim_ids: list[str],
    actor: ActorContext | None = None,
) -> list[dict[str, Any]]:
    """Resolve approved private claims into frozen, provenance-complete IC evidence entries."""
    deal = _deal(session, deal_id, actor)
    if not deal.workspace_id:
        raise IntelligenceError(
            "Approved claims require a deal-linked workspace to enter governed evidence",
            status_code=422,
        )
    if len(claim_ids) != len(set(claim_ids)):
        raise IntelligenceError("approved_claim_ids must be unique", status_code=422)
    manifest: list[dict[str, Any]] = []
    for claim_id in claim_ids:
        claim = _claim(session, claim_id, actor)
        if claim.deal_id != deal_id:
            raise IntelligenceError("Every approved claim must belong to the deal", status_code=422)
        latest = _latest_claim_revision(session, claim.logical_claim_id)
        if latest is None or latest.id != claim.id:
            raise IntelligenceConflict(
                f"Claim '{claim.id}' is not the latest revision and cannot enter an IC packet"
            )
        if claim.review_status != "approved":
            raise IntelligenceConflict(f"Claim '{claim.id}' has not been approved")
        approval = session.scalar(
            select(ClaimReviewEvent).where(
                ClaimReviewEvent.to_claim_id == claim.id,
                ClaimReviewEvent.action == "approve",
            )
        )
        if approval is None or not approval.reviewer_actor_id:
            raise IntelligenceConflict(f"Claim '{claim.id}' lacks an authenticated approval event")
        document, chunk, quoted = _resolved_claim_source(session, claim)
        governed_evidence = _promote_approved_claim(
            session, deal, claim, approval, document, chunk, quoted
        )
        if governed_evidence is None:  # guarded above; keeps the invariant explicit
            raise IntelligenceConflict(f"Claim '{claim.id}' was not promoted to governed evidence")
        entry = {
            "kind": "approved_private_claim",
            "claim_id": claim.id,
            "logical_claim_id": claim.logical_claim_id,
            "revision": claim.revision,
            "category": claim.category,
            "field_name": claim.field_name,
            "value_text": claim.value_text,
            "value_number": claim.value_number,
            "unit": claim.unit,
            "period": claim.period,
            "currency": claim.currency,
            "confidence": claim.confidence,
            "approval": {
                "review_event_id": approval.id,
                "reviewer_actor_id": approval.reviewer_actor_id,
                "approved_at": _utc_iso(approval.created_at),
            },
            "governed_evidence": {
                "evidence_id": governed_evidence.id,
                "ref": governed_evidence.ref,
                "workspace_id": governed_evidence.workspace_id,
                "claim": governed_evidence.claim,
                "claim_type": governed_evidence.claim_type,
                "source_name": governed_evidence.source_name,
                "source_type": governed_evidence.source_type,
                "source_url": governed_evidence.source_url,
                "source_date": governed_evidence.source_date,
                "source_section": governed_evidence.source_section,
                "evidence_text": governed_evidence.evidence_text,
                "confidence": governed_evidence.confidence,
                "agent_name": governed_evidence.agent_name,
            },
            "source": {
                "document_id": document.id,
                "logical_document_id": document.logical_document_id,
                "document_version": document.version,
                "filename": document.filename,
                "document_sha256": document.sha256,
                "chunk_id": chunk.id,
                "chunk_hash": chunk.content_hash,
                "locator": deepcopy(chunk.locator),
                "span": deepcopy(claim.source_span),
            },
        }
        entry["manifest_hash"] = _sha256_text(
            json.dumps(entry, sort_keys=True, separators=(",", ":"), default=str)
        )
        manifest.append(entry)
    return manifest


def list_current_claims(
    session: Session,
    deal_id: str,
    actor: ActorContext | None = None,
) -> dict[str, list[StructuredClaim]]:
    _deal(session, deal_id, actor)
    claims = list(
        session.scalars(
            select(StructuredClaim)
            .where(StructuredClaim.deal_id == deal_id)
            .order_by(StructuredClaim.logical_claim_id, StructuredClaim.revision.desc())
        )
    )
    latest: dict[str, StructuredClaim] = {}
    for claim in claims:
        latest.setdefault(claim.logical_claim_id, claim)
    grouped: dict[str, list[StructuredClaim]] = {
        "approved": [],
        "pending": [],
        "rejected": [],
    }
    for claim in latest.values():
        bucket = "pending" if claim.review_status == "unreviewed" else claim.review_status
        grouped[bucket].append(claim)
    for values in grouped.values():
        values.sort(key=lambda item: (item.category, item.field_name, item.created_at))
    return grouped


def claim_history(
    session: Session,
    logical_claim_id: str,
    actor: ActorContext | None = None,
) -> tuple[list[StructuredClaim], list[ClaimReviewEvent]]:
    revisions = list(
        session.scalars(
            select(StructuredClaim)
            .where(StructuredClaim.logical_claim_id == logical_claim_id)
            .order_by(StructuredClaim.revision)
        )
    )
    if not revisions:
        raise IntelligenceNotFound(f"Logical claim '{logical_claim_id}' not found")
    _deal(session, revisions[0].deal_id, actor)
    reviews = list(
        session.scalars(
            select(ClaimReviewEvent)
            .where(ClaimReviewEvent.logical_claim_id == logical_claim_id)
            .order_by(ClaimReviewEvent.to_revision)
        )
    )
    return revisions, reviews


def _locator_key(chunk: DataRoomChunk) -> str:
    locator = {key: value for key, value in chunk.locator.items() if key not in {"char_start", "char_end"}}
    return json.dumps(locator, sort_keys=True, separators=(",", ":")) + f":{chunk.ordinal}"


def _change_findings(
    from_document: DataRoomDocument,
    to_document: DataRoomDocument,
    before_chunks: list[DataRoomChunk],
    after_chunks: list[DataRoomChunk],
) -> list[dict[str, Any]]:
    before = {_locator_key(chunk): chunk for chunk in before_chunks}
    after = {_locator_key(chunk): chunk for chunk in after_chunks}
    findings: list[dict[str, Any]] = []
    for key in sorted(set(before) | set(after)):
        old = before.get(key)
        new = after.get(key)
        if old and new and old.content_hash == new.content_hash:
            continue
        if old and new:
            finding_type = "modified"
            summary = "Content changed at the same source locator."
        elif old:
            finding_type = "removed"
            summary = "Content was removed in the later version."
        else:
            finding_type = "added"
            summary = "Content was added in the later version."
        findings.append(
            {
                "finding_type": finding_type,
                "summary": summary,
                "before": _citation(from_document, old, old.text) if old else None,
                "after": _citation(to_document, new, new.text) if new else None,
                "shared_terms": sorted(_tokens(old.text) & _tokens(new.text))[:20]
                if old and new
                else [],
            }
        )
        if len(findings) >= 500:
            break
    return findings


def _numbers(text: str) -> set[str]:
    return {match.group(0).strip().casefold() for match in _NUMBER_RE.finditer(text)}


def _contradiction_findings(
    from_document: DataRoomDocument,
    to_document: DataRoomDocument,
    before_chunks: list[DataRoomChunk],
    after_chunks: list[DataRoomChunk],
) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    for old in before_chunks:
        old_terms = _tokens(old.text)
        old_numbers = _numbers(old.text)
        old_negated = bool(re.search(r"\b(?:not|no|never|without)\b", old.text, re.IGNORECASE))
        for new in after_chunks:
            shared = old_terms & _tokens(new.text)
            if len(shared) < 2:
                continue
            new_numbers = _numbers(new.text)
            new_negated = bool(re.search(r"\b(?:not|no|never|without)\b", new.text, re.IGNORECASE))
            numeric_conflict = bool(old_numbers and new_numbers and old_numbers != new_numbers)
            polarity_conflict = old_negated != new_negated
            if not numeric_conflict and not polarity_conflict:
                continue
            conflict_type = "numeric_contradiction" if numeric_conflict else "polarity_contradiction"
            findings.append(
                {
                    "finding_type": conflict_type,
                    "summary": (
                        "Related statements contain different numeric assertions."
                        if numeric_conflict
                        else "Related statements have opposing polarity."
                    ),
                    "before": _citation(from_document, old, old.text),
                    "after": _citation(to_document, new, new.text),
                    "shared_terms": sorted(shared)[:20],
                }
            )
            if len(findings) >= 500:
                return findings
    return findings


def compare_documents(
    session: Session,
    deal_id: str,
    data: ComparisonRequest,
    actor: ActorContext | None = None,
) -> DocumentComparison:
    _deal(session, deal_id, actor)
    from_document = _document(session, data.from_document_id, actor)
    to_document = _document(session, data.to_document_id, actor)
    if from_document.deal_id != deal_id or to_document.deal_id != deal_id:
        raise IntelligenceError("Both comparison documents must belong to the deal", status_code=422)
    if data.comparison_type == "change" and (
        from_document.logical_document_id != to_document.logical_document_id
    ):
        raise IntelligenceError(
            "Change detection requires versions of the same logical document", status_code=422
        )
    before_chunks = list_chunks(session, from_document.id, actor)
    after_chunks = list_chunks(session, to_document.id, actor)
    if data.comparison_type == "change":
        findings = _change_findings(
            from_document, to_document, before_chunks, after_chunks
        )
    else:
        findings = _contradiction_findings(
            from_document, to_document, before_chunks, after_chunks
        )
    comparison = DocumentComparison(
        deal_id=deal_id,
        from_document_id=from_document.id,
        to_document_id=to_document.id,
        comparison_type=data.comparison_type,
        findings=findings,
        finding_count=len(findings),
        algorithm_version=COMPARISON_VERSION,
        created_by_actor_id=_actor_id(actor),
    )
    session.add(comparison)
    return _commit(session, comparison)


def list_comparisons(
    session: Session,
    deal_id: str,
    actor: ActorContext | None = None,
    *,
    limit: int = 100,
) -> list[DocumentComparison]:
    _deal(session, deal_id, actor)
    return list(
        session.scalars(
            select(DocumentComparison)
            .where(DocumentComparison.deal_id == deal_id)
            .order_by(DocumentComparison.created_at.desc())
            .limit(min(max(limit, 1), 1_000))
        )
    )


def _workspace(
    session: Session, workspace_id: str, actor: ActorContext | None = None
) -> Workspace:
    workspace = session.get(Workspace, workspace_id)
    if workspace is None:
        raise IntelligenceNotFound(f"Workspace '{workspace_id}' not found")
    linked_deal = session.scalar(select(Deal).where(Deal.workspace_id == workspace_id))
    if linked_deal:
        _verify_scope(actor, linked_deal.organization_id)
    elif actor and actor.organization_id:
        raise IntelligenceError(
            "A scoped actor can only compare SEC filings for a workspace linked to its deal",
            status_code=403,
        )
    return workspace


def _filing_citation(filing: Filing, chunk: FilingChunk) -> dict[str, Any]:
    return {
        "document_id": filing.id,
        "logical_document_id": filing.accession_number or filing.id,
        "document_version": 1,
        "filename": f"{filing.form_type} filed {filing.filing_date}",
        "sha256": _sha256_text(
            f"{filing.accession_number}|{filing.document_url}|{filing.filing_date}"
        ),
        "chunk_id": chunk.id,
        "content_hash": _sha256_text(chunk.chunk_text),
        "locator": {
            "type": "sec_filing",
            "form_type": filing.form_type,
            "filing_date": filing.filing_date,
            "accession_number": filing.accession_number,
            "section": chunk.section,
            "chunk_index": chunk.chunk_index,
            "source_url": chunk.source_url or filing.document_url,
        },
        "quote": chunk.chunk_text,
    }


def compare_sec_filings(
    session: Session,
    workspace_id: str,
    data: SecFilingComparisonRequest,
    actor: ActorContext | None = None,
) -> SecFilingComparison:
    _workspace(session, workspace_id, actor)
    before_filing = session.get(Filing, data.from_filing_id)
    after_filing = session.get(Filing, data.to_filing_id)
    if before_filing is None or after_filing is None:
        raise IntelligenceNotFound("One or both SEC filings were not found")
    if (
        before_filing.workspace_id != workspace_id
        or after_filing.workspace_id != workspace_id
    ):
        raise IntelligenceError("Both SEC filings must belong to the workspace", status_code=422)
    before_chunks = list(
        session.scalars(
            select(FilingChunk)
            .where(FilingChunk.filing_id == before_filing.id)
            .order_by(FilingChunk.section, FilingChunk.chunk_index)
        )
    )
    after_chunks = list(
        session.scalars(
            select(FilingChunk)
            .where(FilingChunk.filing_id == after_filing.id)
            .order_by(FilingChunk.section, FilingChunk.chunk_index)
        )
    )
    before = {(item.section, item.chunk_index): item for item in before_chunks}
    after = {(item.section, item.chunk_index): item for item in after_chunks}
    findings: list[dict[str, Any]] = []
    for key in sorted(set(before) | set(after)):
        old = before.get(key)
        new = after.get(key)
        if old and new and _normalized(old.chunk_text) == _normalized(new.chunk_text):
            continue
        if old and new:
            finding_type = "modified"
            summary = "SEC disclosure text changed in the corresponding section chunk."
        elif old:
            finding_type = "removed"
            summary = "SEC disclosure text was removed in the later filing."
        else:
            finding_type = "added"
            summary = "SEC disclosure text was added in the later filing."
        findings.append(
            {
                "finding_type": finding_type,
                "summary": summary,
                "before": _filing_citation(before_filing, old) if old else None,
                "after": _filing_citation(after_filing, new) if new else None,
                "shared_terms": sorted(_tokens(old.chunk_text) & _tokens(new.chunk_text))[:20]
                if old and new
                else [],
            }
        )
        if len(findings) >= 500:
            break
    comparison = SecFilingComparison(
        workspace_id=workspace_id,
        from_filing_id=before_filing.id,
        to_filing_id=after_filing.id,
        findings=findings,
        finding_count=len(findings),
        algorithm_version=COMPARISON_VERSION,
        created_by_actor_id=_actor_id(actor),
    )
    session.add(comparison)
    return _commit(session, comparison)


def list_sec_filing_comparisons(
    session: Session,
    workspace_id: str,
    actor: ActorContext | None = None,
    *,
    limit: int = 100,
) -> list[SecFilingComparison]:
    _workspace(session, workspace_id, actor)
    return list(
        session.scalars(
            select(SecFilingComparison)
            .where(SecFilingComparison.workspace_id == workspace_id)
            .order_by(SecFilingComparison.created_at.desc())
            .limit(min(max(limit, 1), 1_000))
        )
    )


def _citation_resolves(session: Session, citation: dict[str, Any]) -> bool:
    document = session.get(DataRoomDocument, citation.get("document_id"))
    chunk = session.get(DataRoomChunk, citation.get("chunk_id"))
    return bool(
        document
        and chunk
        and chunk.document_id == document.id
        and document.sha256 == citation.get("sha256")
        and chunk.content_hash == citation.get("content_hash")
        and citation.get("quote", "") in chunk.text
        and citation.get("locator") == chunk.locator
    )


def _numeric_tokens(text: str) -> set[str]:
    return {match.group("number").replace(",", "") for match in _NUMBER_RE.finditer(text)}


def run_evaluation(
    session: Session,
    deal_id: str,
    data: EvaluationRequest,
    actor: ActorContext | None = None,
) -> IntelligenceEvaluation:
    _deal(session, deal_id, actor)
    runs: list[CitedQARun] = []
    abstention_results: list[bool] = []
    citation_results: list[bool] = []
    numeric_results: list[bool] = []
    content_results: list[bool] = []
    for case in data.cases:
        run = answer_question(
            session,
            deal_id,
            CitedQARequest(question=case.question, filters=case.filters),
            actor,
        )
        runs.append(run)
        abstention_results.append((run.status == "abstained") == case.should_abstain)
        citations_ok = bool(run.citations) if run.status == "answered" else not run.citations
        citations_ok = citations_ok and all(
            _citation_resolves(session, citation) for citation in run.citations
        )
        citation_results.append(citations_ok)
        answer_numbers = _numeric_tokens(run.answer)
        cited_numbers = _numeric_tokens(" ".join(item["quote"] for item in run.citations))
        numeric_results.append(answer_numbers <= cited_numbers)
        normalized_answer = run.answer.casefold()
        content_results.append(
            all(expected.casefold() in normalized_answer for expected in case.expected_answer_contains)
        )

    def ratio(results: list[bool]) -> float:
        return round(sum(results) / len(results), 6) if results else 1.0

    metrics = {
        "case_count": len(data.cases),
        "numeric_traceability": ratio(numeric_results),
        "citation_resolution": ratio(citation_results),
        "abstention_accuracy": ratio(abstention_results),
        "expected_content_accuracy": ratio(content_results),
        "thresholds": {
            "numeric_traceability": data.minimum_numeric_traceability,
            "citation_resolution": data.minimum_citation_resolution,
            "abstention_accuracy": data.minimum_abstention_accuracy,
        },
    }
    passed = bool(
        metrics["numeric_traceability"] >= data.minimum_numeric_traceability
        and metrics["citation_resolution"] >= data.minimum_citation_resolution
        and metrics["abstention_accuracy"] >= data.minimum_abstention_accuracy
        and metrics["expected_content_accuracy"] == 1.0
    )
    evaluation = IntelligenceEvaluation(
        deal_id=deal_id,
        cases=[case.model_dump(mode="json") for case in data.cases],
        qa_run_ids=[run.id for run in runs],
        metrics=metrics,
        passed=passed,
        algorithm_version=EVALUATION_VERSION,
        created_by_actor_id=_actor_id(actor),
    )
    session.add(evaluation)
    return _commit(session, evaluation)


def list_evaluations(
    session: Session,
    deal_id: str,
    actor: ActorContext | None = None,
    *,
    limit: int = 100,
) -> list[IntelligenceEvaluation]:
    _deal(session, deal_id, actor)
    return list(
        session.scalars(
            select(IntelligenceEvaluation)
            .where(IntelligenceEvaluation.deal_id == deal_id)
            .order_by(IntelligenceEvaluation.created_at.desc())
            .limit(min(max(limit, 1), 1_000))
        )
    )
